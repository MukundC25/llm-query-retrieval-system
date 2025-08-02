"""
Simplified LLM-Powered Intelligent Query-Retrieval System
Main FastAPI application - streamlined for hackathon submission
"""

from fastapi import FastAPI, HTTPException, Depends, status, Request
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel, HttpUrl, field_validator
from typing import List, Optional
import os
import logging
import time
import asyncio
import requests
import io
import re
import hashlib
from dotenv import load_dotenv

# Document processing imports
from pdfminer.high_level import extract_text
from pdfminer.layout import LAParams
from docx import Document

# Keep it simple - no additional dependencies

# AI imports
import google.generativeai as genai

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Templates
templates = Jinja2Templates(directory="templates")

app = FastAPI(
    title="LLM Query-Retrieval System",
    description="AI system for processing documents and answering natural language queries",
    version="1.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# Initialize pre-indexing in background (will happen on first request)
# Note: Pre-indexing will be triggered on first API call to avoid startup delays

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security
security = HTTPBearer()

# Expected bearer token
EXPECTED_TOKEN = "12776c804e23764323a141d7736af662e2e2d41a9deaf12e331188a32e1c299f"

def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Verify bearer token authentication"""
    if credentials.credentials != EXPECTED_TOKEN:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid authentication token",
            headers={"WWW-Authenticate": "Bearer"},
        )
    return credentials.credentials

# Request/Response models
class QueryRequest(BaseModel):
    documents: HttpUrl
    questions: List[str]

    @field_validator('questions')
    @classmethod
    def validate_questions(cls, v):
        if not v:
            raise ValueError('At least one question is required')
        if len(v) > 20:
            raise ValueError('Maximum 20 questions allowed per request')
        return v

class QueryResponse(BaseModel):
    answers: List[str]

# Initialize Gemini
def initialize_gemini():
    """Initialize Gemini AI"""
    try:
        api_key = os.getenv('GEMINI_API_KEY')
        if not api_key:
            logger.error("GEMINI_API_KEY environment variable is required")
            return None

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        logger.info("Gemini model initialized successfully")
        return model
    except Exception as e:
        logger.error(f"Failed to initialize Gemini: {e}")
        return None

# Global model instance
gemini_model = None

def get_gemini_model():
    """Get or initialize Gemini model"""
    global gemini_model
    if gemini_model is None:
        gemini_model = initialize_gemini()
    return gemini_model

def generate_embeddings(texts: List[str]) -> List[List[float]]:
    """Generate embeddings using Gemini's embedding API"""
    try:
        embeddings = []
        for text in texts:
            # Use Gemini's embedding model
            result = genai.embed_content(
                model="models/embedding-001",
                content=text,
                task_type="retrieval_document"
            )
            embeddings.append(result['embedding'])
        return embeddings
    except Exception as e:
        logger.error(f"Error generating embeddings: {e}")
        # Fallback to simple embeddings if Gemini fails
        return generate_simple_embeddings(texts)

def generate_simple_embeddings(texts: List[str]) -> List[List[float]]:
    """Fallback simple embeddings"""
    embeddings = []
    for text in texts:
        # Create a simple embedding based on text characteristics
        text_lower = text.lower()
        features = [
            len(text) / 1000.0,
            text.count(' ') / 100.0,
            text.count('.') / 10.0,
            text.count('?') / 5.0,
            text.count('!') / 5.0,
        ]

        # Add hash-based features
        text_hash = hashlib.md5(text_lower.encode()).hexdigest()
        for i in range(0, len(text_hash), 2):
            hex_val = int(text_hash[i:i+2], 16)
            features.append((hex_val - 128) / 128.0)

        # Pad to 768 dimensions (Gemini embedding size)
        while len(features) < 768:
            features.append(0.0)
        features = features[:768]
        embeddings.append(features)

    return embeddings

def cosine_similarity_manual(vec1: List[float], vec2: List[float]) -> float:
    """Calculate cosine similarity without external dependencies"""
    dot_product = sum(a * b for a, b in zip(vec1, vec2))
    magnitude1 = sum(a * a for a in vec1) ** 0.5
    magnitude2 = sum(b * b for b in vec2) ** 0.5

    if magnitude1 == 0 or magnitude2 == 0:
        return 0.0

    return dot_product / (magnitude1 * magnitude2)

# Document processing functions
async def download_document(url: str) -> bytes:
    """Download document from URL"""
    try:
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        return response.content
    except requests.RequestException as e:
        logger.error(f"Error downloading document from {url}: {e}")
        raise Exception(f"Failed to download document: {e}")

def extract_text_from_pdf(pdf_content: bytes) -> str:
    """Extract text from PDF content using PDFMiner with optimized settings"""
    try:
        # Use optimized LAParams for better text extraction
        laparams = LAParams(
            boxes_flow=0.5,
            word_margin=0.1,
            char_margin=2.0,
            line_margin=0.5
        )

        text = extract_text(
            io.BytesIO(pdf_content),
            laparams=laparams
        )

        logger.info(f"Extracted {len(text)} characters from PDF")
        return text

    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise Exception(f"Failed to extract PDF text: {e}")

def extract_text_from_docx(docx_content: bytes) -> str:
    """Extract text from DOCX content"""
    try:
        doc = Document(io.BytesIO(docx_content))
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return '\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise Exception(f"Failed to extract DOCX text: {e}")

def clean_text(text: str) -> str:
    """Enhanced text cleaning and normalization for better processing"""
    # Remove excessive whitespace and normalize line breaks
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)

    # Fix common PDF extraction issues
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between camelCase
    text = re.sub(r'(\d)([A-Za-z])', r'\1 \2', text)  # Space between numbers and letters
    text = re.sub(r'([A-Za-z])(\d)', r'\1 \2', text)  # Space between letters and numbers

    # Normalize insurance-specific terms
    insurance_normalizations = {
        r'waitingperiod': 'waiting period',
        r'graceperiod': 'grace period',
        r'suminsured': 'sum insured',
        r'roomrent': 'room rent',
        r'preexisting': 'pre-existing',
        r'co-payment': 'copayment',
        r'co payment': 'copayment'
    }

    for pattern, replacement in insurance_normalizations.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

    # Remove special characters but keep important punctuation
    text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\/\%\$]', ' ', text)

    # Normalize spacing around punctuation
    text = re.sub(r'\s*([\.,:;!?])\s*', r'\1 ', text)

    # Remove extra spaces
    text = re.sub(r'\s+', ' ', text)

    return text.strip()

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[dict]:
    """Split text into overlapping chunks with smart boundary detection"""
    chunks = []
    text_length = len(text)
    start = 0
    chunk_id = 0

    while start < text_length:
        end = min(start + chunk_size, text_length)

        # Try to break at sentence boundary for better semantic coherence
        if end < text_length:
            # Look for sentence endings within the last 200 characters
            for punct in ['. ', '.\n', '! ', '!\n', '? ', '?\n']:
                sentence_end = text.rfind(punct, max(start, end - 200), end)
                if sentence_end > start:
                    end = sentence_end + len(punct)
                    break

        chunk_text = text[start:end].strip()

        if chunk_text and len(chunk_text) > 50:  # Only include meaningful chunks
            chunks.append({
                'id': chunk_id,
                'text': chunk_text,
                'start_pos': start,
                'end_pos': end,
                'length': len(chunk_text)
            })
            chunk_id += 1

        # Move start position with overlap
        start = end - overlap
        if start >= text_length:
            break

    return chunks

async def process_document(document_url: str) -> dict:
    """Process document from URL with embeddings"""
    try:
        # Download document
        logger.info(f"Downloading document from: {document_url}")
        content = await download_document(document_url)

        # Determine file type and extract text
        if str(document_url).lower().endswith('.pdf') or 'pdf' in str(document_url).lower():
            text = extract_text_from_pdf(content)
        elif str(document_url).lower().endswith('.docx'):
            text = extract_text_from_docx(content)
        else:
            # Try PDF first, then DOCX
            try:
                text = extract_text_from_pdf(content)
            except:
                text = extract_text_from_docx(content)

        # Clean and chunk text
        cleaned_text = clean_text(text)
        chunks = chunk_text(cleaned_text)

        # Generate embeddings for all chunks
        logger.info(f"Generating embeddings for {len(chunks)} chunks...")
        chunk_texts = [chunk['text'] for chunk in chunks]
        embeddings = generate_embeddings(chunk_texts)

        # Add embeddings to chunks
        for i, chunk in enumerate(chunks):
            chunk['embedding'] = embeddings[i]

        logger.info(f"Processed document: {len(chunks)} chunks, {len(cleaned_text)} characters")

        return {
            'url': document_url,
            'full_text': cleaned_text,
            'chunks': chunks,
            'total_chunks': len(chunks),
            'total_characters': len(cleaned_text),
            'embeddings': embeddings
        }

    except Exception as e:
        logger.error(f"Error processing document: {e}")
        raise Exception(f"Document processing failed: {e}")

def find_relevant_chunks(query: str, document_data: dict, top_k: int = 5) -> List[dict]:
    """Find most relevant chunks using enhanced semantic similarity with reranking"""
    try:
        # Generate embedding for the query
        query_embedding = generate_embeddings([query])[0]
        query_lower = query.lower()

        # Calculate similarities with all chunks
        similarities = []
        for chunk in document_data['chunks']:
            if 'embedding' in chunk:
                # Base semantic similarity
                semantic_similarity = cosine_similarity_manual(query_embedding, chunk['embedding'])

                # Enhanced scoring with multiple factors
                chunk_text_lower = chunk['text'].lower()

                # Keyword overlap bonus
                query_words = set(query_lower.split())
                chunk_words = set(chunk_text_lower.split())
                keyword_overlap = len(query_words.intersection(chunk_words)) / max(len(query_words), 1)

                # Length penalty for very short chunks
                length_factor = min(1.0, len(chunk['text']) / 100)

                # Insurance domain relevance bonus
                domain_keywords = ['policy', 'coverage', 'benefit', 'claim', 'premium', 'insured',
                                 'waiting', 'grace', 'exclusion', 'maternity', 'ayush', 'room rent']
                domain_score = sum(1 for keyword in domain_keywords if keyword in chunk_text_lower) / len(domain_keywords)

                # Combined score
                final_score = (
                    semantic_similarity * 0.6 +
                    keyword_overlap * 0.25 +
                    domain_score * 0.1 +
                    length_factor * 0.05
                )

                similarities.append({
                    'chunk': chunk,
                    'similarity': final_score,
                    'semantic_sim': semantic_similarity,
                    'keyword_overlap': keyword_overlap
                })

        # Sort by enhanced similarity and return top_k
        similarities.sort(key=lambda x: x['similarity'], reverse=True)

        relevant_chunks = []
        for item in similarities[:top_k]:
            chunk_data = item['chunk'].copy()
            chunk_data['similarity_score'] = item['similarity']
            chunk_data['semantic_score'] = item['semantic_sim']
            relevant_chunks.append(chunk_data)

        logger.info(f"Found {len(relevant_chunks)} relevant chunks with enhanced scores: {[f'{c['similarity_score']:.3f}' for c in relevant_chunks]}")
        return relevant_chunks

    except Exception as e:
        logger.error(f"Error finding relevant chunks: {e}")
        return []

def find_text_matches(question: str, document_data: dict) -> List[dict]:
    """Find chunks containing exact text matches for key terms with enhanced pattern matching"""
    question_lower = question.lower()

    # Enhanced term mapping with variations and synonyms
    term_patterns = {
        'waiting period': ['waiting period', 'wait period', 'waiting time', 'wait time', 'moratorium'],
        'grace period': ['grace period', 'grace time', 'premium grace', 'payment grace'],
        'premium': ['premium', 'premium payment', 'premium amount', 'premium due'],
        'maternity': ['maternity', 'pregnancy', 'childbirth', 'delivery', 'natal'],
        'sum insured': ['sum insured', 'sum assured', 'coverage amount', 'insured amount', 'maximum benefit'],
        'room rent': ['room rent', 'room charges', 'accommodation', 'hospital room'],
        'cataract': ['cataract', 'eye surgery', 'lens replacement'],
        'ayush': ['ayush', 'ayurvedic', 'homeopathic', 'unani', 'siddha', 'alternative medicine'],
        'exclusion': ['exclusion', 'excluded', 'not covered', 'limitation', 'restriction'],
        'organ donor': ['organ donor', 'donor', 'transplant', 'organ donation'],
        'policy term': ['policy term', 'policy period', 'term of policy', 'policy duration'],
        'pre-existing': ['pre-existing', 'pre existing', 'existing disease', 'prior condition'],
        'coverage': ['coverage', 'cover', 'benefit', 'protection'],
        'claim': ['claim', 'reimbursement', 'settlement'],
        'deductible': ['deductible', 'co-payment', 'copayment', 'co pay']
    }

    # Find relevant patterns from question
    relevant_patterns = []
    for main_term, variations in term_patterns.items():
        for variation in variations:
            if variation in question_lower:
                relevant_patterns.extend(variations)
                break

    # If no specific patterns found, extract key words
    if not relevant_patterns:
        question_words = question_lower.split()
        relevant_patterns = [word for word in question_words if len(word) > 3]

    # Find chunks containing these patterns
    matching_chunks = []
    for chunk in document_data['chunks']:
        chunk_text_lower = chunk['text'].lower()
        match_score = 0
        matched_terms = []

        for pattern in relevant_patterns:
            if pattern in chunk_text_lower:
                match_score += 1
                matched_terms.append(pattern)

        if match_score > 0:
            # Calculate relevance score based on number of matches
            relevance_score = min(1.0, match_score / len(relevant_patterns))
            matching_chunks.append({
                'chunk': chunk,
                'matched_terms': matched_terms,
                'text_score': relevance_score,
                'match_count': match_score
            })

    # Sort by relevance score
    matching_chunks.sort(key=lambda x: x['text_score'], reverse=True)

    return matching_chunks

async def answer_question_with_context(question: str, relevant_chunks: List[dict], document_url: str = "") -> str:
    """Generate answer using Gemini AI with relevant context chunks and caching"""
    try:
        # Check cache first to save API quota
        if document_url:
            cached_answer = get_cached_response(document_url, question)
            if cached_answer:
                return cached_answer

        model = get_gemini_model()

        if model is None:
            return "Error: AI service is not available. Please try again later."

        if not relevant_chunks:
            return "This information is not available in the provided document."

        # Smart context preparation with merging and optimization
        context_parts = []
        total_length = 0
        max_context_length = 4000  # Keep within token limits

        # Sort chunks by relevance and merge adjacent ones if beneficial
        sorted_chunks = sorted(relevant_chunks, key=lambda x: x.get('similarity_score', 0), reverse=True)

        for i, chunk in enumerate(sorted_chunks):
            similarity = chunk.get('similarity_score', chunk.get('text_score', 0))
            chunk_text = chunk['text']

            # Skip very low relevance chunks unless we have few chunks
            if similarity < 0.1 and len(sorted_chunks) > 3:
                continue

            # Check if adding this chunk would exceed context limit
            if total_length + len(chunk_text) > max_context_length and len(context_parts) > 0:
                break

            context_parts.append(f"[Document Section {i+1} - Relevance: {similarity:.2f}]\n{chunk_text}\n")
            total_length += len(chunk_text)

        context = "\n".join(context_parts)
        logger.info(f"Prepared context with {len(context_parts)} sections, total length: {total_length} chars")

        # Enhanced prompt with better context management and accuracy focus
        prompt = f"""You are an expert insurance document analyst. Your task is to provide accurate, specific answers based ONLY on the provided document sections.

DOCUMENT CONTEXT:
{context}

USER QUESTION: {question}

ANALYSIS FRAMEWORK:
1. SCAN ALL SECTIONS: Carefully examine each document section for relevant information
2. IDENTIFY KEY TERMS: Look for specific terms related to the question
3. EXTRACT PRECISE DETAILS: Find exact numbers, percentages, conditions, and requirements
4. CROSS-REFERENCE: Check if information spans multiple sections for complete context

RESPONSE GUIDELINES:
✓ PROVIDE SPECIFIC DETAILS: Include exact figures, time periods, conditions
✓ QUOTE RELEVANT TEXT: Use direct quotes when providing specific information
✓ REFERENCE SECTIONS: Mention which section contains the information
✓ BE COMPREHENSIVE: If information is found, provide complete details
✓ STAY FACTUAL: Only use information explicitly stated in the document

SEARCH PRIORITIES FOR INSURANCE QUERIES:
- Waiting periods → "waiting period", "wait", "moratorium", "cooling period"
- Grace periods → "grace period", "grace time", "premium grace", "payment grace"
- Coverage limits → "sum insured", "coverage amount", "maximum benefit", "limit"
- Exclusions → "excluded", "not covered", "limitation", "restriction", "exception"
- Benefits → "covered", "benefit", "included", "eligible"
- Conditions → "subject to", "provided that", "conditions", "requirements"

IMPORTANT: If the specific information is not found in any document section, respond: "This information is not available in the provided document."

ANSWER:"""

        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                max_output_tokens=1200,
                temperature=0.05  # Even lower temperature for more consistency
            )
        )

        if response and response.text:
            answer = response.text.strip()

            # Add source reference if answer contains information
            if "not available" not in answer.lower() and len(relevant_chunks) > 0:
                best_similarity = relevant_chunks[0].get('similarity_score', relevant_chunks[0].get('text_score', 0))
                if best_similarity > 0.2:  # Lower threshold for text matches
                    answer += f"\n\n[Source: Policy document section with {best_similarity:.1%} relevance]"

            # Cache the response for future use
            if document_url and "error" not in answer.lower():
                cache_response(document_url, question, answer)

            return answer
        else:
            return "Error: No response generated from AI service"

    except Exception as e:
        logger.error(f"Error generating answer: {e}")
        return f"Error: Unable to generate answer - {str(e)}"

# Document cache and response cache
document_cache = {}
response_cache = {}  # Cache for question-answer pairs to optimize API usage

def get_cache_key(document_url: str, question: str) -> str:
    """Generate cache key for question-document combination"""
    combined = f"{document_url}:{question.lower().strip()}"
    return hashlib.md5(combined.encode()).hexdigest()

def cache_response(document_url: str, question: str, answer: str):
    """Cache response for future use"""
    cache_key = get_cache_key(document_url, question)
    response_cache[cache_key] = {
        'answer': answer,
        'timestamp': time.time()
    }

    # Keep cache size manageable (max 100 entries)
    if len(response_cache) > 100:
        # Remove oldest entries
        oldest_keys = sorted(response_cache.keys(),
                           key=lambda k: response_cache[k]['timestamp'])[:20]
        for key in oldest_keys:
            del response_cache[key]

def get_cached_response(document_url: str, question: str) -> str:
    """Get cached response if available"""
    cache_key = get_cache_key(document_url, question)
    if cache_key in response_cache:
        cached = response_cache[cache_key]
        # Cache valid for 1 hour
        if time.time() - cached['timestamp'] < 3600:
            logger.info(f"Using cached response for question: {question[:50]}...")
            return cached['answer']
        else:
            # Remove expired cache
            del response_cache[cache_key]
    return None

async def process_queries(document_url: str, questions: List[str]) -> List[str]:
    """Process all queries for a document using hybrid search"""
    try:
        # Check cache first
        if document_url in document_cache:
            logger.info("Using cached document data")
            document_data = document_cache[document_url]
        else:
            # Process document
            logger.info("Processing new document")
            document_data = await process_document(document_url)
            document_cache[document_url] = document_data

        # Answer all questions using hybrid search
        answers = []
        for i, question in enumerate(questions):
            try:
                logger.info(f"Processing question {i+1}/{len(questions)}: {question}")

                # Enhanced chunk selection with better ranking
                semantic_chunks = find_relevant_chunks(question, document_data, top_k=4)

                # Find chunks with text pattern matching
                text_match_chunks = find_text_matches(question, document_data)

                # Smart chunk combination and ranking
                all_chunks = []
                chunk_ids_seen = set()

                # Add semantic chunks with enhanced scoring
                for chunk_data in semantic_chunks:
                    chunk_id = chunk_data.get('id', id(chunk_data))
                    if chunk_id not in chunk_ids_seen:
                        all_chunks.append(chunk_data)
                        chunk_ids_seen.add(chunk_id)

                # Add text match chunks with their scores
                for match_data in text_match_chunks:
                    chunk = match_data['chunk']
                    chunk_id = chunk.get('id', id(chunk))
                    if chunk_id not in chunk_ids_seen:
                        chunk_with_score = chunk.copy()
                        chunk_with_score['text_score'] = match_data['text_score']
                        all_chunks.append(chunk_with_score)
                        chunk_ids_seen.add(chunk_id)

                # Sort by combined relevance and select top chunks
                def get_chunk_score(chunk):
                    semantic_score = chunk.get('similarity_score', 0)
                    text_score = chunk.get('text_score', 0)
                    return max(semantic_score, text_score)  # Take the best score

                all_chunks.sort(key=get_chunk_score, reverse=True)
                relevant_chunks = all_chunks[:4]  # Limit to top 4 for better context

                logger.info(f"Selected {len(relevant_chunks)} top-ranked chunks")

                # Generate answer with caching
                answer = await answer_question_with_context(question, relevant_chunks, document_url)
                answers.append(answer)

                logger.info(f"Generated answer for Q{i+1}: {answer[:100]}...")

            except Exception as e:
                logger.error(f"Error processing question '{question}': {e}")
                answers.append(f"Error: Unable to process this question - {str(e)}")

        return answers

    except Exception as e:
        logger.error(f"Error in process_queries: {e}")
        return [f"Error: Document processing failed - {str(e)}"] * len(questions)

# Routes
@app.get("/", response_class=HTMLResponse)
async def frontend(request: Request):
    """Frontend interface for the LLM Query Retrieval System"""
    return templates.TemplateResponse("index.html", {"request": request})

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "message": "LLM Query Retrieval System is running",
        "timestamp": time.time(),
        "version": "1.0.0"
    }

@app.get("/api/v1/hackrx/run")
async def hackrx_webhook_verification():
    """HackRX webhook endpoint - GET method for webhook verification"""
    return {
        "status": "success",
        "message": "HackRX webhook endpoint is active",
        "project_name": "LLM Query Retrieval System",
        "team": "AI Innovators",
        "hackrx_submission": True,
        "webhook_verified": True,
        "timestamp": time.time()
    }

@app.post("/api/v1/hackrx/run", response_model=QueryResponse)
async def hackrx_document_processing(
    request: QueryRequest,
    req: Request,
    token: str = Depends(verify_token)
):
    """Main endpoint for processing document queries"""
    start_time = time.time()

    try:
        logger.info(f"Processing request with {len(request.questions)} questions for document: {request.documents}")

        # Process the document and questions with timeout
        try:
            answers = await asyncio.wait_for(
                process_queries(
                    document_url=str(request.documents),
                    questions=request.questions
                ),
                timeout=25.0  # 25 second timeout
            )

            processing_time = time.time() - start_time
            logger.info(f"Request processed successfully in {processing_time:.2f} seconds")

        except asyncio.TimeoutError:
            processing_time = time.time() - start_time
            logger.warning(f"Query processing timed out after {processing_time:.2f} seconds")
            timeout_message = "Processing timed out. The document may be too large or complex."
            answers = [timeout_message] * len(request.questions)

        return QueryResponse(answers=answers)

    except Exception as e:
        processing_time = time.time() - start_time
        logger.error(f"Error processing queries after {processing_time:.2f} seconds: {e}")
        error_message = f"Error processing queries: {str(e)}"
        error_answers = [error_message] * len(request.questions)
        return QueryResponse(answers=error_answers)

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
