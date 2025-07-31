"""
Simplified LLM-Powered Intelligent Query-Retrieval System
Main FastAPI application - streamlined for hackathon submission
"""

from fastapi import FastAPI, HTTPException, Depends, status, Request
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel, HttpUrl, field_validator
from typing import List, Optional
import os
import logging
import time
import asyncio
import requests
import io
import re
from dotenv import load_dotenv

# Document processing imports
from pdfminer.high_level import extract_text
from pdfminer.layout import LAParams
from docx import Document

# AI imports
import google.generativeai as genai

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Templates
templates = Jinja2Templates(directory="templates")

app = FastAPI(
    title="LLM Query-Retrieval System",
    description="AI system for processing documents and answering natural language queries",
    version="1.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security
security = HTTPBearer()

# Expected bearer token
EXPECTED_TOKEN = "12776c804e23764323a141d7736af662e2e2d41a9deaf12e331188a32e1c299f"

def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Verify bearer token authentication"""
    if credentials.credentials != EXPECTED_TOKEN:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid authentication token",
            headers={"WWW-Authenticate": "Bearer"},
        )
    return credentials.credentials

# Request/Response models
class QueryRequest(BaseModel):
    documents: HttpUrl
    questions: List[str]

    @field_validator('questions')
    @classmethod
    def validate_questions(cls, v):
        if not v:
            raise ValueError('At least one question is required')
        if len(v) > 20:
            raise ValueError('Maximum 20 questions allowed per request')
        return v

class QueryResponse(BaseModel):
    answers: List[str]

# Initialize Gemini
def initialize_gemini():
    """Initialize Gemini AI"""
    api_key = os.getenv('GEMINI_API_KEY')
    if not api_key:
        raise ValueError("GEMINI_API_KEY environment variable is required")
    
    genai.configure(api_key=api_key)
    return genai.GenerativeModel('gemini-1.5-flash')

# Global model instance
gemini_model = None

def get_gemini_model():
    """Get or initialize Gemini model"""
    global gemini_model
    if gemini_model is None:
        gemini_model = initialize_gemini()
    return gemini_model

# Document processing functions
async def download_document(url: str) -> bytes:
    """Download document from URL"""
    try:
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        return response.content
    except requests.RequestException as e:
        logger.error(f"Error downloading document from {url}: {e}")
        raise Exception(f"Failed to download document: {e}")

def extract_text_from_pdf(pdf_content: bytes) -> str:
    """Extract text from PDF content"""
    try:
        laparams = LAParams(
            boxes_flow=0.5,
            word_margin=0.1,
            char_margin=2.0,
            line_margin=0.5
        )
        
        text = extract_text(
            io.BytesIO(pdf_content),
            laparams=laparams
        )
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise Exception(f"Failed to extract PDF text: {e}")

def extract_text_from_docx(docx_content: bytes) -> str:
    """Extract text from DOCX content"""
    try:
        doc = Document(io.BytesIO(docx_content))
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return '\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise Exception(f"Failed to extract DOCX text: {e}")

def clean_text(text: str) -> str:
    """Clean extracted text"""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', ' ', text)
    return text.strip()

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[dict]:
    """Split text into overlapping chunks"""
    chunks = []
    start = 0
    chunk_id = 0
    
    while start < len(text):
        end = start + chunk_size
        chunk_text = text[start:end]
        
        chunks.append({
            'id': chunk_id,
            'text': chunk_text,
            'start_pos': start,
            'end_pos': end,
            'length': len(chunk_text)
        })
        
        start += chunk_size - overlap
        chunk_id += 1
    
    return chunks

async def process_document(document_url: str) -> dict:
    """Process document from URL"""
    try:
        # Download document
        logger.info(f"Downloading document from: {document_url}")
        content = await download_document(document_url)
        
        # Determine file type and extract text
        if str(document_url).lower().endswith('.pdf') or 'pdf' in str(document_url).lower():
            text = extract_text_from_pdf(content)
        elif str(document_url).lower().endswith('.docx'):
            text = extract_text_from_docx(content)
        else:
            # Try PDF first, then DOCX
            try:
                text = extract_text_from_pdf(content)
            except:
                text = extract_text_from_docx(content)
        
        # Clean and chunk text
        cleaned_text = clean_text(text)
        chunks = chunk_text(cleaned_text)
        
        logger.info(f"Processed document: {len(chunks)} chunks, {len(cleaned_text)} characters")
        
        return {
            'url': document_url,
            'full_text': cleaned_text,
            'chunks': chunks,
            'total_chunks': len(chunks),
            'total_characters': len(cleaned_text)
        }
        
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        raise Exception(f"Document processing failed: {e}")

async def answer_question(question: str, document_text: str) -> str:
    """Generate answer using Gemini AI"""
    try:
        model = get_gemini_model()
        
        # Create prompt for question answering
        prompt = f"""
You are an expert document analyst. Based on the provided document content, answer the following question accurately and concisely.

Document Content:
{document_text[:4000]}  # Limit context to avoid token limits

Question: {question}

Instructions:
1. Provide a direct, accurate answer based only on the document content
2. If the information is not in the document, say "This information is not available in the provided document"
3. Include specific details, numbers, or conditions when available
4. Keep the answer concise but complete

Answer:"""

        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                max_output_tokens=800,
                temperature=0.1
            )
        )
        
        return response.text.strip()
        
    except Exception as e:
        logger.error(f"Error generating answer: {e}")
        return f"Error: Unable to generate answer - {str(e)}"

# Document cache
document_cache = {}

async def process_queries(document_url: str, questions: List[str]) -> List[str]:
    """Process all queries for a document"""
    try:
        # Check cache first
        if document_url in document_cache:
            logger.info("Using cached document data")
            document_data = document_cache[document_url]
        else:
            # Process document
            logger.info("Processing new document")
            document_data = await process_document(document_url)
            document_cache[document_url] = document_data
        
        # Answer all questions
        answers = []
        for question in questions:
            try:
                answer = await answer_question(question, document_data['full_text'])
                answers.append(answer)
            except Exception as e:
                logger.error(f"Error processing question '{question}': {e}")
                answers.append(f"Error: Unable to process this question - {str(e)}")
        
        return answers
        
    except Exception as e:
        logger.error(f"Error in process_queries: {e}")
        return [f"Error: Document processing failed - {str(e)}"] * len(questions)

# Routes
@app.get("/", response_class=HTMLResponse)
async def frontend(request: Request):
    """Frontend interface for the LLM Query Retrieval System"""
    return templates.TemplateResponse("index.html", {"request": request})

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "message": "LLM Query Retrieval System is running",
        "timestamp": time.time(),
        "version": "1.0.0"
    }

@app.get("/api/v1/hackrx/run")
async def hackrx_webhook_verification():
    """HackRX webhook endpoint - GET method for webhook verification"""
    return {
        "status": "success",
        "message": "HackRX webhook endpoint is active",
        "project_name": "LLM Query Retrieval System",
        "team": "AI Innovators",
        "hackrx_submission": True,
        "webhook_verified": True,
        "timestamp": time.time()
    }

@app.post("/api/v1/hackrx/run", response_model=QueryResponse)
async def hackrx_document_processing(
    request: QueryRequest,
    req: Request,
    token: str = Depends(verify_token)
):
    """Main endpoint for processing document queries"""
    start_time = time.time()

    try:
        logger.info(f"Processing request with {len(request.questions)} questions for document: {request.documents}")

        # Process the document and questions with timeout
        try:
            answers = await asyncio.wait_for(
                process_queries(
                    document_url=str(request.documents),
                    questions=request.questions
                ),
                timeout=25.0  # 25 second timeout
            )

            processing_time = time.time() - start_time
            logger.info(f"Request processed successfully in {processing_time:.2f} seconds")

        except asyncio.TimeoutError:
            processing_time = time.time() - start_time
            logger.warning(f"Query processing timed out after {processing_time:.2f} seconds")
            timeout_message = "Processing timed out. The document may be too large or complex."
            answers = [timeout_message] * len(request.questions)

        return QueryResponse(answers=answers)

    except Exception as e:
        processing_time = time.time() - start_time
        logger.error(f"Error processing queries after {processing_time:.2f} seconds: {e}")
        error_message = f"Error processing queries: {str(e)}"
        error_answers = [error_message] * len(request.questions)
        return QueryResponse(answers=error_answers)

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
