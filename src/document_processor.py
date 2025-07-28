"""
Document Processing Module
Handles document ingestion, parsing, and text extraction
"""

import requests
import io
import re
from typing import List, Dict, Any
from pdfminer.high_level import extract_text
from pdfminer.layout import LAParams
from docx import Document
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document processing from various sources and formats"""
    
    def __init__(self):
        self.chunk_size = 1000  # Characters per chunk
        self.chunk_overlap = 200  # Overlap between chunks
    
    async def download_document(self, url: str) -> bytes:
        """Download document from URL"""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            return response.content
        except requests.RequestException as e:
            logger.error(f"Error downloading document from {url}: {e}")
            raise Exception(f"Failed to download document: {e}")
    
    def extract_text_from_pdf(self, pdf_content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            # Use LAParams for better text extraction
            laparams = LAParams(
                boxes_flow=0.5,
                word_margin=0.1,
                char_margin=2.0,
                line_margin=0.5
            )
            
            text = extract_text(
                io.BytesIO(pdf_content),
                laparams=laparams
            )
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise Exception(f"Failed to extract PDF text: {e}")
    
    def extract_text_from_docx(self, docx_content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            doc = Document(io.BytesIO(docx_content))
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise Exception(f"Failed to extract DOCX text: {e}")
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\"\'\/]', '', text)
        
        # Remove extra newlines
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    def chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks for better retrieval"""
        chunks = []
        text_length = len(text)
        
        start = 0
        chunk_id = 0
        
        while start < text_length:
            # Calculate end position
            end = min(start + self.chunk_size, text_length)
            
            # Try to break at sentence boundary
            if end < text_length:
                # Look for sentence endings within the last 100 characters
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + self.chunk_size - 100:
                    end = sentence_end + 1
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunks.append({
                    'id': chunk_id,
                    'text': chunk_text,
                    'start_pos': start,
                    'end_pos': end,
                    'length': len(chunk_text)
                })
                chunk_id += 1
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            if start >= text_length:
                break
        
        return chunks
    
    async def process_document(self, document_url: str) -> Dict[str, Any]:
        """Main method to process document from URL"""
        try:
            # Download document
            logger.info(f"Downloading document from: {document_url}")
            content = await self.download_document(document_url)
            
            # Determine file type from URL or content
            if document_url.lower().endswith('.pdf'):
                text = self.extract_text_from_pdf(content)
            elif document_url.lower().endswith('.docx'):
                text = self.extract_text_from_docx(content)
            else:
                # Try PDF first, then DOCX
                try:
                    text = self.extract_text_from_pdf(content)
                except:
                    text = self.extract_text_from_docx(content)
            
            # Clean and chunk text
            cleaned_text = self.clean_text(text)
            chunks = self.chunk_text(cleaned_text)
            
            logger.info(f"Processed document: {len(chunks)} chunks, {len(cleaned_text)} characters")
            
            return {
                'url': document_url,
                'full_text': cleaned_text,
                'chunks': chunks,
                'total_chunks': len(chunks),
                'total_characters': len(cleaned_text)
            }
            
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            raise Exception(f"Document processing failed: {e}")
